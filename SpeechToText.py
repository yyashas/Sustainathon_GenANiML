import speech_recognition as sr
from docx import Document
import threading
from io import BytesIO
from scipy.io.wavfile import write
import numpy as np

class SpeechToText:
    """
    A class to perform speech-to-text conversion using a microphone input or an audio file.
    
    Attributes:
    -----------
    language : str
        The language code for speech recognition.
    stop_listening : bool
        Flag to control the listening loop.
    recognized_text_storage : str
        Storage for recognized text for saving into a document.
    
    Methods:
    --------
    recognize_speech():
        Captures audio from the microphone and returns the recognized text.
    recognize_speech_from_file(audio_file):
        Processes audio from a file-like object and returns the recognized text.
    continuous_speech_to_text():
        Continuously performs speech recognition until stopped.
    start_listening():
        Starts the speech-to-text conversion in a separate thread.
    stop_listening_func():
        Stops the speech-to-text conversion.
    save_to_word(file_name="output"):
        Saves the recognized text to a Word document.
    get_recognized_text():
        Returns the recognized text.
    """
    
    def __init__(self, language='en-US'):
        """Initialize the speech recognizer and default settings."""
        self.recognizer = sr.Recognizer()
        self.stop_listening = False
        self.recognized_text_storage = ""
        self.language = language  # Set the language based on input

    def recognize_speech(self):
        """
        Captures audio from the microphone and returns recognized text.
        
        Returns:
        --------
        str : Recognized text or error message.
        """
        mic = sr.Microphone()

        try:
            with mic as source:
                self.recognizer.adjust_for_ambient_noise(source, duration=1)
                self.recognizer.energy_threshold = 300
                self.recognizer.pause_threshold = 0.8

                print("Listening...")
                audio = self.recognizer.listen(source, phrase_time_limit=8)

            text = self.recognizer.recognize_google(audio, language=self.language)
            return text

        except sr.RequestError:
            return "[Error: Unable to reach the speech recognition service]"
        except sr.UnknownValueError:
            return "[Error: Could not understand audio]"
        except Exception as e:
            return f"[Error: {str(e)}]"

    def recognize_speech_from_file(self, audio_file):
        """
        Processes audio from a file-like object and returns recognized text.
        
        Parameters:
        -----------
        audio_file : file-like object
            The audio file to transcribe.
        
        Returns:
        --------
        str : Recognized text or error message.
        """
        try:
            with sr.AudioFile(audio_file) as source:
                audio = self.recognizer.record(source)
                text = self.recognizer.recognize_google(audio, language=self.language)
                return text
        except sr.RequestError:
            return "[Error: Unable to reach the speech recognition service]"
        except sr.UnknownValueError:
            return "[Error: Could not understand audio]"
        except Exception as e:
            return f"[Error: {str(e)}]"

    def continuous_speech_to_text(self):
        """Continuously performs speech recognition until stopped."""
        while not self.stop_listening:
            recognized_text = self.recognize_speech()
            if recognized_text and not recognized_text.startswith("[Error"):
                print(f"Recognized: {recognized_text}")
                self.recognized_text_storage += recognized_text + " "

    def start_listening(self):
        """Starts the speech-to-text conversion in a separate thread."""
        self.stop_listening = False
        print("Started listening...")
        threading.Thread(target=self.continuous_speech_to_text, daemon=True).start()

    def stop_listening_func(self):
        """Stops the speech-to-text conversion."""
        self.stop_listening = True
        print("Stopped listening.")

    def save_to_word(self, file_name="output"):
        """
        Saves the recognized text to a Word document.
        
        Parameters:
        -----------
        file_name : str
            The name of the file to save the document as.
        """
        if self.recognized_text_storage.strip():
            doc = Document()
            doc.add_heading(f'{self.language} Speech to Text', 0)
            doc.add_paragraph(self.recognized_text_storage.strip())
            doc.save(f"{file_name}.docx")
            print(f"Document saved as {file_name}.docx")
        else:
            print("No text to save!")

    def get_recognized_text(self):
        """
        Returns the recognized text.
        
        Returns:
        --------
        str : The recognized text stored during the session.
        """
        return self.recognized_text_storage.strip()

# Example usage
if __name__ == "__main__":
    # Choose language
    lang_choice = input("Choose language: '1' for English, '2' for Kannada\n").strip()

    if lang_choice == '1':
        language_code = 'en-US'
    elif lang_choice == '2':
        language_code = 'kn-IN'
    else:
        print("Invalid choice, defaulting to English.")
        language_code = 'en-US'

    stt = SpeechToText(language=language_code)
    audio_data = np.zeros((44100 * 5,), dtype=np.int16)  # Example audio data
    audio_buffer = BytesIO()
    write(audio_buffer, 44100, audio_data)  # Example: write as WAV format
    audio_buffer.seek(0)

    print(stt.recognize_speech_from_file(audio_buffer))